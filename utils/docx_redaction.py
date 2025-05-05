from typing import List, Dict, Any
from docx import Document
import re
from io import BytesIO
from utils.ner import NERProcessor

# Initialize NER processor
ner_processor = NERProcessor()

class RedactionRule:
    def __init__(self, type_: str, value: str, name: str, rule_id: str, is_ai_detected: bool):
        self.type = type_
        self.value = value
        self.name = name
        self.rule_id = rule_id
        self.is_ai_detected = is_ai_detected

def getSpacyText(text, value):
    """
    Extract entities from text using NER processor.
    
    Args:
        text (str): The text to process
        value (str or list): The entity type(s) to extract
        
    Returns:
        list: List of extracted entities
    """
    try:
        entities = ner_processor.extract_entities(text)
        # Filter entities by type if specified
        if value:
            # Convert single value to list for consistent handling
            entity_types = [value] if isinstance(value, str) else value
            # Filter entities by any of the specified types
            entities = [ent for ent in entities if ent['type'].lower() in [t.lower() for t in entity_types]]
        return [ent['text'] for ent in entities]
    except Exception as e:
        print(f"Error in getSpacyText: {e}")
        return []

def redact_docx(document_bytes: bytes, rules: List[RedactionRule], template_id: str) -> Dict[str, Any]:
    doc = Document(BytesIO(document_bytes))
    report = {"redactions": [], "template_id": template_id}
    total_redactions = 0
    # Extract initial text for before_text
    initial_text = ""
    for para in doc.paragraphs:
        initial_text += para.text + "\n"
    report["before_text"] = initial_text
    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text
        new_text = para_text
        para_report = []
        for rule in rules:
            if rule.type == 'text':
                matches = [(m.start(), m.end(), m.group()) for m in re.finditer(re.escape(rule.value), para_text)]
            elif rule.type == 'regex':
                matches = [(m.start(), m.end(), m.group()) for m in re.finditer(rule.value, para_text)]
            elif rule.type == 'spacy':
                matches = []
                # Get all matching entities
                entities = getSpacyText(initial_text, rule.value)
                # Create matches for each entity
                for entity in entities:
                    # Find all occurrences of this entity in the text
                    entity_matches = [(m.start(), m.end(), m.group()) for m in re.finditer(re.escape(entity), para_text)]
                    matches.extend(entity_matches)
            else:
                continue
            for start, end, match_text in matches:
                # Replace match with black box (█)
                replacement = '█' * (end - start)
                new_text = new_text.replace(match_text, replacement, 1)
                para_report.append({
                    "rule": rule.name,
                    "type": rule.type,
                    "text": match_text,
                    "paragraph": para_idx + 1,
                    "rule_id": rule.rule_id,
                    "index":total_redactions,
                    "is_ai_detected": rule.is_ai_detected
                })
                total_redactions += 1
        if para_report:
            report["redactions"].extend(para_report)
        if new_text != para_text:
            para.text = new_text
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    # Extract final text for after_text
    final_text = ""
    doc_final = Document(BytesIO(out.read()))
    for para in doc_final.paragraphs:
        final_text += para.text + "\n"
    report["after_text"] = final_text
    # Format before_text in markdown with yellow background highlights
    before_text_md = initial_text
    for redaction in report["redactions"]:
        before_text_md = before_text_md.replace(redaction["text"], f"<span style='background-color: yellow;'>{redaction['text']}</span>")
    report["before_text"] = before_text_md
    report["total_redactions"] = total_redactions
    return {"redacted_docx": out.read(), "report": report} 